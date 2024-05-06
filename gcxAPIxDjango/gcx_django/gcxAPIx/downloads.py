import re
import docx
import html
import time
import json
import pdfkit
import datetime
import requests
from random import randint
from pyvirtualdisplay import Display
from rest_framework import status
from django.http import HttpResponse
from rest_framework.response import Response
from rest_framework.decorators import api_view
from .factory import getConnection


# from flask_app.server import loglogger
# from flask import Blueprint, Response, request, send_file

# doc_dload_bp = Blueprint('doc_download_bp', __name__)


@api_view(['GET'])
def downloadPDFByDocId(requset, doc_id, *args, **kwargs):
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT contents, title FROM user_documents WHERE doc_url= '%s' """ % doc_id)
    rows = cur.fetchall()
    try:
        contents = rows[0][0]
    except IndexError:
        return Response({'response': 'File does not exist'}, status.HTTP_204_NO_CONTENT)
    # contents = contents.replace("'", "#x27;")
    if rows[0][1] and rows[0][1] != '':
        title = rows[0][1]
        for ch in ['#', '<', '$', '+', '%', '>', '!', '`', '&', '*', "'", '|', '{', '?', '\"', '=', '}', '/', ':', '\\',
                   '@']:
            if ch in title:
                title = title.replace(ch, '')
        title = title.replace(' ', '-')
    else:
        title = "Untitled-Document"
    # path = '/var/www/GCX/docx/' + title + '-' + doc_id + '.pdf'
    try:
        display = Display(visible=False, size=(1366, 768))
        display.start()
        options = {
            'page-size': 'A4',
            'margin-top': '0.25in',
            'margin-right': '0.25in',
            'margin-bottom': '0.25in',
            'margin-left': '0.25in',
            'encoding': "UTF-8",
            'quiet': '',
            'no-outline': None
        }
        # body ='<h2>'+title+'</h2>\n\n<p>'+contents+'</p>'
        body = '<meta http-equiv="Content-Type" content="text/html;charset=utf-8" /><p>' + contents + '</p>'
        pdf_new = pdfkit.from_string(body, options=options)
        file_name = title + '-' + doc_id + '.pdf'
        response = HttpResponse(pdf_new, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="' + file_name + '"'
    except Exception as e:
        print(e)
        return Response({"response": "Error in execution"}, status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        display.stop()
    conn.close()
    return response
    # return send_file(path, as_attachment=True)


################################################################################################


@api_view(['GET'])
def downloadDOCByDocId(requset, doc_id, *args, **kwargs):
    # print(doc_id)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT title, contents  FROM user_documents WHERE doc_url='%s' """ % doc_id)
    rows = cur.fetchall()
    try:
        if rows[0][0] and rows[0][0] != '':
            title = rows[0][0]
            for ch in ['#', '<', '$', '+', '%', '>', '!', '`', '&', '*', "'", '|', '{', '?', '\"', '=', '}', '/', ':', '\\',
                       '@']:
                if ch in title:
                    title = title.replace(ch, '')
            title = title.replace(' ', '-')
        else:
            title = "Untitled-Document"
        contents = rows[0][1]
    except IndexError:
        return Response({'response': 'Document does not exist'}, status.HTTP_204_NO_CONTENT)
    ##################################################
    # create an instance of a word document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    # doc.add_heading(title, 0)
    # add a paragraph and store
    # the object in a variable
    doc.add_paragraph(contents)
    # now save the document to a location
    doc.save('/var/www/GCX/docx/' + title + '-' + doc_id + '.docx')
    path = '/var/www/GCX/docx/' + title + '-' + doc_id + '.docx'
    try:
        with open(path, 'rb') as file:
            response = HttpResponse(file, content_type='application/application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            file_name = title + '-' + doc_id + '.docx'
            response['Content-Disposition'] = 'attachment; filename="' + file_name + '"'
        #file.close()
        # file = f.read()
    except Exception as e:
        print(e)
        return Response({"response": "Error in execution"}, status.HTTP_500_INTERNAL_SERVER_ERROR)
    conn.close()
    return response
    # return send_file(path, as_attachment=True)


################################################################################################
# Google Drive Download by Muneeb
################################################################################################


@api_view(['POST'])
# @auth_token_required
def downloadFromDrive(request, *args, **kwargs):
    if 'profile_id' in request.data and 'file_path' in request.data and 'file_type' in request.data and 'hash' in request.data:
        user_id = request.data['profile_id']
        file_path = request.data['file_path']
        file_type = request.data['file_type']
        token = request.data['hash']
        create_chunks = request.data.get('create_chunks', 'false')
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status.HTTP_206_PARTIAL_CONTENT)
    randName = str(randint(1, 99999999999999))
    headersg = {"Authorization": "Bearer " + token}
    if file_type == 'docx':
        sampleFileName = "/var/www/GCX/docx/" + randName + ".docx"
        urlg = "https://www.googleapis.com/drive/v3/files/" + file_path + \
               "/export?mimeType=application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        files = requests.request("GET", urlg, headers=headersg)
        f = open(sampleFileName, "wb")
        f.write(files.content)
        f.close()
        try:
            doc = docx.Document(sampleFileName)
        except Exception as e:
            with open(sampleFileName, "r") as f:
                file = json.loads(f.read())
                if file['error']['code'] == 403:
                    urlg = "https://www.googleapis.com/drive/v3/files/" + file_path + "?alt=media"
                    files = requests.request("GET", urlg, headers=headersg)
                    f = open(sampleFileName, "wb")
                    f.write(files.content)
                    f.close()
                    doc = docx.Document(sampleFileName)
                else:
                    print("File Data++>>\n", files.content)
                    print(e)
                    json_response = {"status": 501, "message": 'Error reading file'}
                    code = 501
                    return Response(json_response, status=code)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        contents = '\n'.join(fullText)
    elif file_type == 'txt':
        sampleFileName = "/var/www/GCX/docx/" + randName + ".txt"
        urlg = "https://www.googleapis.com/drive/v3/files/" + file_path + "?alt=media"
        files = requests.request("GET", urlg, headers=headersg)
        f = open(sampleFileName, "wb")
        f.write(files.content)
        f.close()
        fullText = []
        with open(sampleFileName, "r") as f:
            lines = f.readlines()
            for para in lines:
                fullText.append(para)
            contents = '\n'.join(fullText)
    else:
        sampleFileName = "/var/www/GCX/docx/GCX-Report-" + randName + ".pdf"
        urlg = "https://www.googleapis.com/drive/v3/files/" + file_path + "?alt=media"
        files = requests.request("GET", urlg, headers=headersg)
        f = open(sampleFileName, "wb")
        f.write(files.content)
        f.close()
        url = "http://v2.pdfapis.com/api/v2/TextExtractor"
        payload = {'authToken': '7FE9B6118D6484C8314F0378399FBD05', 'type': '3'}
        files = [('', ('.pdf', open(sampleFileName, 'rb'), 'application/pdf'))]
        headers = {}
        response = requests.request("POST", url, headers=headers, data=payload, files=files)
        content = request.get_json(silent=True)
        r = response.json()
        try:
            contents = r['text'].replace('\x00', '')
            contents = re.sub("(\t)+", " ", contents)
        except Exception:
            # loglogger.error("Error Happend in execution.", exc_info=1)
            json_response = json.dumps({"message": 'Error reading file'})
            code = 501
            return Response(json_response, status=code)
    contents = re.sub("\r\n(\s)+", "\n", contents.strip())
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100, 999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    sql = "INSERT INTO user_documents(user_id, doc, contents, doc_url, spell_errors, total_words, grammar_errors, suggestion_errors, score) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id"
    try:
        cur.execute(sql, (user_id, doc, contents, doc_url, 0, 0, 0, 0, 0))
        # Disabbled because creating double entries in database.
        # conn.commit()
        # Chunk Slicing is working and is commented till further notice.
        chunklist = []
        if create_chunks.lower() == 'true':
            from .documents import chunkslicing
            chnk_OBJ = chunkslicing()
            chnk_OBJ.checkConditions({'text': contents, 'length': len(contents), 'offset': 0, 'id': 0, })
            contents, chunklist = chnk_OBJ.makechunks()
    except Exception:
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        conn.close()
        return Response(json_response, status=400)
    new_id = cur.fetchone()[0]
    data = {'new_id': new_id, 'new_doc_id': new_doc_id,
            'new_contents': contents, 'USER_ID': user_id, 'chunk_ids': chunklist}
    conn.close()
    #json_response = json.dumps(data, indent=4, sort_keys=True, default=str)
    return Response(data, status=status.HTTP_200_OK)
