# from django.shortcuts import render
import docx
import json
import time
import pdfkit
import datetime
from pyvirtualdisplay import Display
# from django.http import HttpResponse
from django.contrib.auth import get_user_model
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
#from rest_framework.authtoken.models import Token
from .auth import bearer_auth
from .factory import getConnection
from .functions import sendAttachmentEmailPdf, sendAttachmentEmailWord, initializeFilters
from .serializers import UserSerializer

User = get_user_model()


# Create your views here.

class UserViewSet(viewsets.ModelViewSet):
    """
    API endpoint that allows users to be viewed or edited.
    """
    queryset = User.objects.all().order_by('-date_joined')
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]


@api_view(['GET', 'POST'])
def signUpUser1(*args, **kwargs):
    print("Server active and Responding...")
    data = {'contents': '123'}
    return Response(data, status=status.HTTP_200_OK)


@api_view(['POST'])
@bearer_auth
def activateUserAccount(*args, **kwargs):
    """Get data returned from the server."""
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("UPDATE auth_user SET first_name = 'Admin1' WHERE id = 1;")
    conn.commit()
    cur.execute("SELECT * FROM auth_user;")
    rows = cur.fetchall()
    data = {'contents': rows[0][5]}
    conn.close()
    return Response(data, status=status.HTTP_200_OK)


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getUserProfileContents(request, *args, **kwargs):
    """ Doc """
    if 'profile_id' in request.data:
        profile_id = request.data['profile_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    # """SELECT meta_value FROM user_meta WHERE meta_name=%s AND user_id=%s """, (aurgmentName, profile_id)
    cur.execute(
        """SELECT email, first_name, last_name, dict_status, login_from FROM auth_user WHERE id=%s """ % profile_id)
    rows = cur.fetchall()
    data = {'getUserProfileContents': '1', 'contents': rows}
    conn.close()
    return Response(data, status=status.HTTP_200_OK)


@api_view(['POST'])
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def submitEDocument(request, *args, **kwargs):
    """Get data returned from the server."""
    if 'profile_id' in request.data and 'doc_url' in request.data and 'email_sub' in request.data and 'email_body' in request.data:
        profile_id = request.data['profile_id']
        doc_url = request.data['doc_url']
        email_sub = request.data['email_sub']
        email_body = request.data['email_body']
        attachment_type = request.data['attachment_type']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    # edit by Muneeb
    cur.execute("""SELECT email FROM auth_user WHERE id = '%s'; """ %
                profile_id)
    email_to = cur.fetchall()[0][0]
    # email_cc = ''
    cur.execute(
        """SELECT title, contents FROM user_documents WHERE doc_url='%s' """ % doc_url)
    rows = cur.fetchall()
    # print(rows)
    if rows[0][0] and rows[0][0] != '':
        title = rows[0][0]
        for ch in ['#', '<', '$', '+', '%', '>', '!', '`', '&', '*', "'", '|', '{', '?', '\"', '=', '}', '/', ':', '\\',
                   '@']:
            if ch in title:
                title = title.replace(ch, '')
        title = title.replace(' ', '-')
    else:
        title = "Untitled-Document"
    contents = rows[0][1]
    ##################################################
    if attachment_type == 'pdf':
        path = '/var/www/GCX/docx/' + title + '-' + doc_url + '.pdf'
        display = Display(visible=False, size=(1366, 768))
        display.start()
        try:
            options = {
                'page-size': 'A4',
                'margin-top': '0.25in',
                'margin-right': '0.25in',
                'margin-bottom': '0.25in',
                'margin-left': '0.25in',
                'encoding': "UTF-8",
                'quiet': '',
                'no-outline': None
            }
            # body ='<h2>'+title+'</h2>\n\n<p>'+contents+'</p>'
            body = '<meta http-equiv="Content-Type" content="text/html;charset=utf-8" /><p>' + contents + '</p>'
            pdfkit.from_string(body, path, options=options)

            from_address = 'gcx@grammarin.com'
            # if email_cc == '':
            output = sendAttachmentEmailPdf(
                from_address, email_to, email_sub, email_body, path)
            # else:
            #    output = sendAttachmentCCEmailPdf(
            #        from_address, email_to, email_cc, email_sub, email_body, path, doc_url)
            # pdfkit.from_url('https://www.grammarin.com/#/share/156353516920622', '/var/www/GCX/docx/out1.pdf')
        finally:
            display.stop()
        # send_file(path, as_attachment=True)    # commented because of no use
    else:
        # create an instance of a Word document
        doc = docx.Document()
        # add a heading of level 0 (the largest heading)
        # doc.add_heading(title, 0)
        # add a paragraph and store
        # the object in a variable
        doc.add_paragraph(contents)
        # now save the document to a location
        doc.save('/var/www/GCX/docx/' + title + '-' + doc_url + '.docx')
        path = '/var/www/GCX/docx/' + title + '-' + doc_url + '.docx'
        from_address = 'gcx@grammarin.com'
        # if email_cc == '':
        output = sendAttachmentEmailWord(
            from_address, email_to, email_sub, email_body, path)
        # else:
        #    output = sendAttachmentCCEmailWord(
        #        from_address, email_to, email_cc, email_sub, email_body, path, doc_url)
    conn.close()
    return Response(output, status=200)


@api_view(['GET', 'POST'])
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def helpCentre(request, *args, **kwargs):
    """ Help """
    conn = getConnection()
    cur = conn.cursor()
    if 'firstname' in request.data and 'email' in request.data and 'subject' in request.data and 'reference_no' in request.data and 'message' in request.data:
        firstname = request.data['firstname']
        email = request.data['email']
        subject = request.data['subject']
        reference = request.data['reference_no']
        message = request.data['message']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    try:
        cur.execute(
            "INSERT INTO help_faqs(first_name, email, subject, reference_no, message) VALUES(%s, %s, %s, %s, %s) RETURNING id;",
            (firstname, email, subject, reference, message))
        user_id = cur.fetchone()[0]
        cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                    (user_id, 'preferred_language', 'us'))
        conn.commit()
    except Exception as e:
        # print("Error Happend in execution.")
        print(e)
        conn.rollback()
        json_response = 'Feedback submission failed'
        print(json_response)
        conn.close()
        return Response({'response': json_response}, status=400)
    # FIRST WE WILL VERFY IF THE USER ALREADY THERE, IF YES, THEN JUST LOG HIM IN.
    # cur.execute("""SELECT * FROM help_faqs WHERE email = '%s' """ % email)
    # result = cur.fetchall()
    # if len(result) > 0:
    data = {
        'help': '1', 'help_message': 'Thank you for your feedback. We will get back to you soon'}
    conn.close()
    return Response(data, status=status.HTTP_200_OK)


# this API request was not used and is disabled from URL.py
@api_view(['GET', 'POST'])
@bearer_auth
def getHelpFAQsListing(*args, **kwargs):
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(""" SELECT * FROM help_faqs WHERE status = 'active' """)
    rows = cur.fetchall()
    data = {'getHelpFAQsListing': '1', 'contents': rows}
    conn.close()
    return Response(data, status=200)


################################################################################################

@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def userWizard(request, *args, **kwargs):
    # The following statements are calculating execution time
    # from datetime import datetime
    # start_time = datetime.now()
    if 'id' in request.data and 'skip_intro' in request.data:
        id = request.data['id']
        skip_intro = request.data['skip_intro']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""SELECT login_count FROM auth_user WHERE id = '%s' """ % id)
    result = cur.fetchall()
    login_count = result[0][0]
    if login_count < 2 and skip_intro == '':
        code = 200
        response = dict()
        response['login_wizard'] = 1
        response['login_wizard_message'] = 'User is eligible for login wizard. Please take him into login wizard'
        response['resp_code'] = 'wzrd_1'
    else:
        if skip_intro == '1':
            try:
                cur.execute(
                    "UPDATE auth_user SET login_count = '2' WHERE id = '%s' " % id)
                conn.commit()
            except Exception as e:
                print("Error Happend in execution.", e)
                conn.rollback()
                json_response = 'user wizard skip failed'
                conn.close()
                return Response({'response': json_response}, status=400)
        code = 200
        response = dict()
        response['login_wizard'] = 0
        response['login_wizard_message'] = 'User is not eligible for login wizard. User has already taken login wizard'
        response['resp_code'] = 'wzrd_0'
    conn.close()
    return Response(dict(meta=dict(code=code), response=response))


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getPreferredLangByProfileID(request, *args, **kwargs):
    """ Language """
    if 'profile_id' in request.data:
        user_id = request.data['profile_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT meta_value FROM user_meta WHERE meta_name = 'preferred_language' AND user_id=%s """ % user_id)
    rows = cur.fetchall()
    if len(rows) <= 0:
        data = {'getPreferredLangByProfileID': '1', 'contents': 'us'}
    else:
        data = {'getPreferredLangByProfileID': '1', 'contents': rows[0][0]}
    conn.close()
    return Response(data, status=200)


@api_view(['GET', 'POST'])
@bearer_auth
def addReportedRule(request, *args, **kwargs):
    if 'profile_id' in request.data and 'rule' in request.data and 'doc_id' in request.data:
        user_id = request.data['profile_id']
        rule = request.data['rule']
        doc_id = request.data['doc_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    try:
        cur.execute(
            "INSERT INTO rule_report (user_id, doc_id, rule_id, doc) VALUES('%s', '%s', '%s', '%s') RETURNING id " % (
                user_id, doc_id, rule, doc))
        new_id = cur.fetchone()[0]
        data = {'addReportedRule': '1', 'new_id': new_id}
        conn.commit()
    except Exception as e:
        print("Error Happend in execution.", e)
        conn.rollback()
        json_response = 'Adding rule failed'
        return Response({'response': json_response}, status=400)
    finally:
        conn.close()
    return Response(data, status=200)


@api_view(['POST'])
@bearer_auth
def addCompTexttoDB(request, *args, **kwargs):
    if 'profile_id' in request.data and 'url' in request.data and 'query' in request.data and 'des' in request.data:
        # user_id = request.data['profile_id']
        url = request.data['url']
        # print(url)
        query = request.data['query']
        des = request.data['des']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    # ts = time.time()
    # doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    try:
        cur.execute(
            "INSERT INTO plag_compare_text (url, query, des) VALUES('%s', '%s', '%s') RETURNING id " % (
                url, query, des))
        new_id = cur.fetchone()[0]
        data = {'addCompTexttoDB': '1', 'new_id': new_id}
        conn.commit()
    except Exception as e:
        print("Error Happend in execution.", e)
        conn.rollback()
        json_response = 'Comp text saving failed'
        return Response({'response': json_response}, status=400)
    finally:
        conn.close()
    return Response(data, status=200)


@api_view(['POST'])
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getuserfilters(request, *args, **kwargs):
    try:
        user_id = request.data['user_id']
        docURL = request.data['doc_url']
        filterstatus = request.data['filterstatus'] if ('filterstatus' in request.data) else 'aftersignup'
    except Exception as e:
        print("Error Happend in execution.", e)
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    curr = conn.cursor()
    msg = ""
    try:
        query = ("SELECT filters FROM user_filters where user_id = '%s' AND doc_id = '%s';" % (
            user_id, docURL))
        curr.execute(query)
        row = curr.fetchone()
        if row:
            msg = "Document already Exist"
        else:
            query = (
                    "SELECT filters FROM user_filters where user_id = '%s' AND doc_id = 'account';" % user_id)
            curr.execute(query)
            row = curr.fetchone()
            if row:
                msg = "Account Settings Used"
                results = row[0]
                query = (
                        "INSERT INTO user_filters(user_id, doc_id, filters) VALUES('%s', '%s', '%s') RETURNING filters;"
                        % (user_id, docURL, json.dumps(results)))
                curr.execute(query)
                conn.commit()
            else:
                row = initializeFilters(user_id, docURL, filterstatus)
    except Exception as e:
        conn.rollback()
        print("Error Happend in execution.", e)
        response = {
            'response': 'There was some issue, Please Try again',
            'status': "Failed",
        }
        return Response(response, status=status.HTTP_206_PARTIAL_CONTENT)
    finally:
        conn.close()
    results = row[0]
    response = {
        'filters': results,
        'resp_code': 200,
        'msg': msg,
    }
    return Response(response, status=status.HTTP_200_OK)


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getuserfiltersAccount(request, *args, **kwargs):
    try:
        user_id = request.data['user_id']
        filterstatus = request.data['filterstatus'] if ('filterstatus' in request.data) else 'aftersignup'
    except Exception as e:
        print("Error Happend in execution.", e)
        response = {
            'response': 'Please Provide all parameters.',
            'status': "Failed",
            'code': "400",
        }
        return Response(response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    curr = conn.cursor()
    try:
        query = ("SELECT filters FROM user_filters where user_id = '%s' AND doc_id = '%s';" % (
            user_id, 'account'))
        curr.execute(query)
        row = curr.fetchone()
        if not row:
            row = initializeFilters(user_id, 'account', filterstatus)
    except Exception as e:
        conn.rollback()
        print("Error Happend in execution.", e)
        response = {
            'response': 'There was some issue, Please Try again',
            'status': "Failed",
        }
        return Response(response, status=status.HTTP_304_NOT_MODIFIED)
    finally:
        conn.close()
    results = row[0]
    response = {
        'filters': results, 'resp_code': 200, 'msg': "Success",
    }
    return Response(response, status=status.HTTP_200_OK)


@api_view(['POST'])
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def addIgnoreWord(request, *args, **kwargs):
    try:
        user_id = request.data['user_id']
        docURL = request.data['doc_url']
        word = request.data['word']
        offset = request.data['offset']
    except Exception as e:
        print("Error Happend in execution.", e)
        response = {
            'response': 'Please Provide all parameters.',
            'status': "Failed",
            'code': "400",
        }
        return Response(response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    curr = conn.cursor()
    try:
        query = ("INSERT INTO user_ignore_words(user_id, doc_id, word, position) VALUES ('%s', '%s', '%s', '%s')" % (
            user_id, docURL, word, offset))
        curr.execute(query)
        conn.commit()
    except Exception as e:
        conn.rollback()
        print("Error Happend in execution.", e)
        response = {
            'response': 'Unable to add word',
            'status': "Failed",
        }
        return Response(response, status=status.HTTP_304_NOT_MODIFIED)
    finally:
        conn.close()
    response = {
        'response': 'Word added',
        'status': "Success"
    }
    return Response(response, status=status.HTTP_200_OK)


@api_view(['POST'])
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def ignoreWordList(request, *args, **kwargs):
    try:
        docURL = request.data['doc_url']
    except Exception as e:
        print("Error Happend in execution.", e)
        response = {
            'response': 'Please Provide all parameters.',
            'status': "Failed",
            'code': "400",
        }
        return Response(response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    curr = conn.cursor()
    try:
        query = (
                "SELECT word, position FROM user_ignore_words where doc_id = '%s' AND active = true;" % docURL)
        curr.execute(query)
        wordList = curr.fetchall()
    except Exception as e:
        conn.rollback()
        print("Error Happend in execution.", e)
        response = {
            'response': 'Unable to add word',
            'status': "Failed",
        }
        return Response(response, status=status.HTTP_304_NOT_MODIFIED)
    finally:
        conn.close()
    response = {
        'response': 'List and offset of words',
        'words_list': wordList,
        'status': "Success",
    }
    return Response(response, status=status.HTTP_200_OK)
