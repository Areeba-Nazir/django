import docx
import time
import json
import base64
import datetime
import requests
import operator
import subprocess
import urllib.parse
from random import randint
from flask_app.server import app, loglogger
from flask import Blueprint, Response, request, jsonify
from flask_security import auth_token_required
from flask_app.factory import getConnection, blocked
import flask_app.app_utils.token_login as logintoken

blocked_bp = Blueprint('blocked_bp', __name__)

################################################################################################
# BELOW ARE THE FUNCTIONS AND APIs THAT ARE NOT BIENG USED NOW
################################################################################################


# ---#------------------------------------------------------------------------------------------#
# --- DASHBOARD SECTION #
# ----------------------------------------------------------------------------------------------#
# --- Language :  Form Save & Resend Frsh Data   (US / UK)


@api_view(['POST'])
def breaktext(request, *args, **kwargs):
    start_time_s = time.time()
    if 'txt' in request.data and 'doc_url' in request.data:
        text = request.data['txt']
        doc_id = request.data['doc_url']
        language = 'en-US'
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    # print('==>>\n',text,'\n')
    # print('==>>\n',doc_id,'\n')
    url = "http://34.206.85.223:8089/api/sentenceBreakage"
    user_agent = 'Mozilla/5.0 (Windows NT 6.1; Win64; x64)'
    headers = {'User-Agent': user_agent, 'Connection': 'close', 'Cache-Control': 'max-age=600',
               'Content-Type': 'application/x-www-form-urlencoded'}
    start_time_r = time.time()
    try:
        text = text.encode('ascii', 'ignore')
        values = {'language': language, 'text': text.decode()}
        response = requests.request("POST", url, headers=headers, data=values)
        output = response.json()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        return Response({'response': 'Unable to break content.', 'status': "Failed", 'code': "500"},
                        status.HTTP_500_INTERNAL_SERVER_ERROR)
    end_time_r = time.time()
    total_time_r = end_time_r - start_time_r
    print("  Break Sentence exe time ==>\t", total_time_r)
    chunks = []
    for val in range(0, len(output['matches'])):
        start = output['matches'][val]['offset']
        if val == (len(output['matches']) - 1):
            end = len(text)
        else:
            end = output['matches'][val + 1]['offset']
        chunks.append({
            "chunk_id": "%s-%s" % (doc_id, val),
            "start": start,
            "end": end,
            "text": text[start:end]
        })
    end_time_s = time.time()
    total_time_s = end_time_s - start_time_s
    print("  Exe Time for remaining Code ==>> \t", total_time_s - total_time_r)
    return Response(chunks, status.HTTP_200_OK)


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def saveDataCorrectContents(request, *args, **kwargs):
    """ Language """
    conn = getConnection()
    cur = conn.cursor()
    if 'profile_id' in request.data and 'doc_id' in request.data and 'docTitle' in request.data and 'docContents' in request.data:
        profile_id = request.data['profile_id']
        doc_id = request.data['doc_id']
        docTitle = request.data['docTitle']
        docContents = request.data['docContents']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    # NOW UPDATE IN THE USER-META TABLE AGAINST USER-ID,
    try:
        cur.execute(
            """UPDATE user_documents SET title=%s, contents=%s WHERE doc_url=%s AND user_id=%s """,
            (docTitle, docContents, doc_id, profile_id))
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving Data failed'
        conn.close()
        return Response(json_response, status=400)
    # AND SELECT AGAIN DATA FROM THE DATABASE AND RETURN TO THE USER
    cur.execute(
        """SELECT * FROM user_documents WHERE doc_url='%s' """ % doc_id)
    rows = cur.fetchall()
    # NOW AS WE HAVE STORED THE DATA INTO THE DATABASE.
    # SO NOW WE WILL GET THE METHOD OF THE LANG. CHECK
    language = 'en-US'
    tool = language_check.LanguageTool(language)
    matches = tool.check(docTitle)
    outputTitle = prepareOutputWeb(matches, language)
    matches = tool.check(docContents)
    docContents = prepareOutputWeb(matches, language)
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    # NOW WE UPDATE THE TOTAL ERRORS COMING INTO THE TITLE & CONTENTS.
    try:
        cur.execute(
            """UPDATE user_documents SET spell_errors=%s, total_words=%s, grammar_errors=%s, suggestion_errors=%s, score=%s WHERE doc_url=%s AND user_id=%s """,
            (spell_errors, total_words, grammar_errors, suggestion_errors, score, doc_id, profile_id))
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Updatung Document attributes Failed.'
        return Response(json_response, status=400)
    finally:
        conn.close()
    data = {'saveDataCorrectContents': '1',
            'contents': rows,
            'USER_ID': profile_id,
            'outputTitle': outputTitle,
            'docContents': docContents
            }
    return Response(data, status.HTTP_200_OK)



@blocked_bp.route('/api/signUpUser_BLOCKED', methods=['GET', 'POST'])
@blocked
def signUpUser_BLOCKED():
    """View function for signup view."""
    content = request.get_json(silent=True)
    if 1:
        firstname = base64.b64decode(content['firstname']).decode('utf-8')
        lastname = base64.b64decode(content['lastname']).decode('utf-8')
        email = base64.b64decode(content['email']).decode('utf-8')
        password = base64.b64decode(content['password']).decode('utf-8')
        active = 'TRUE'
        ts = time.time()
        doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
        #cipher_password = encrypt(encdcr_key, password)
        #plain_password = decrypt(encdcr_key, cipher_password)
        cipher_password = password
        # SELECT DATA FROM THE DATABASE
        conn = getConnection()
        cur = conn.cursor()
        cur.execute(
            """SELECT * FROM auth_user WHERE email = '%s' """ % email)
        result = cur.fetchall()
        if len(result) <= 0:
            cur.execute("INSERT INTO auth_user(first_name, last_name, email, password, active, created_at) VALUES(%s, %s, %s, %s, %s, %s) RETURNING id;",
                        (firstname, lastname, email, cipher_password, active, doc))
            conn.commit()
            user_id = cur.fetchone()[0]
            cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                        (user_id, 'preferred_language', 'us'))
            conn.commit()
            code = 200
            response = dict()
            response['signup'] = 1
            response['signup_message'] = 'Account Created Successfully, Please Sign In to Continue'
            response['user_id'] = user_id
        else:
            code = 400
            response = dict()
            data = {
                'signup': '0', 'signup_message': 'Sorry! this account already exists', 'user_id': 0}
            response['errors'] = data
        conn.close()
        return jsonify(dict(meta=dict(code=code), response=response))
#########################################################
# Added by Haseeb
#########################################################


@blocked_bp.route('/api/signupinGoogle_BLOCKED', methods=['GET', 'POST'])
@blocked
def signupinGoogle_BLOCKED():
    """View function for signup view."""
    content = request.get_json(silent=True)
    if 'firstname' in request.form and 'lastname' in request.form and 'email' in request.form:
        #print ("This is if part")

        # actual_firstname = request.form['firstname'].encode('UTF-8')
        # print(actual_firstname)
        # firstname_base64 = base64.b64encode(actual_firstname)
        # print(firstname_base64)
        # print(firstname_base64.decode('UTF-8'))
        # firstname_base64_decoded = base64.decodebytes(firstname_base64)
        # firstname = firstname_base64_decoded.decode('UTF-8')
        # print(firstname)

        actual_firstname = request.form['firstname'].encode('UTF-8')
        print(actual_firstname)
        print(type(actual_firstname))
        firstname_base64 = base64.b64encode(actual_firstname)
        print(firstname_base64)
        str_decoded = firstname_base64.decode()
        print(type(str_decoded))
        print(str_decoded)
        firstname_base64_decoded = base64.b64decode(firstname_base64)
        firstname1 = base64.b64decode(firstname_base64_decoded)
        firstname = firstname1.decode('UTF-8')
        print("First name after decodingbytes: ", firstname)

        actual_lastname = request.form['lastname'].encode('UTF-8')
        print(actual_lastname)
        lastname_base64 = base64.b64encode(actual_lastname)
        print(lastname_base64)
        print(lastname_base64.decode('UTF-8'))
        lastname_base64_decoded = base64.b64decode(lastname_base64)
        lastname1 = base64.b64decode(lastname_base64_decoded)
        lastname = lastname1.decode('UTF-8')
        print("Last name after decodingbytes: ", lastname)

        actual_email = request.form['email'].encode('UTF-8')
        print(actual_email)
        email_base64 = base64.b64encode(actual_email)
        print(email_base64)
        print(email_base64.decode('UTF-8'))
        email_base64_decoded = base64.b64decode(email_base64)
        email1 = base64.b64decode(email_base64_decoded)
        email = email1.decode('UTF-8')
        print("Email name after decodingbytes: ", email)
        password = '123455678'
    else:
        #print ("This is else part")
        content = request.get_json(silent=True)

        actual_firstname = content['firstname'].encode('UTF-8')
        print(actual_firstname)
        print(type(actual_firstname))
        firstname_base64 = base64.b64encode(actual_firstname)
        print(firstname_base64)
        str_decoded = firstname_base64.decode()
        print(type(str_decoded))
        print(str_decoded)
        firstname_base64_decoded = base64.b64decode(firstname_base64)
        firstname1 = base64.b64decode(firstname_base64_decoded)
        firstname = firstname1.decode('UTF-8')
        print("First name after decodingbytes: ", firstname)

        actual_lastname = content['lastname'].encode('UTF-8')
        print(actual_lastname)
        lastname_base64 = base64.b64encode(actual_lastname)
        print(lastname_base64)
        print(lastname_base64.decode('UTF-8'))
        lastname_base64_decoded = base64.b64decode(lastname_base64)
        lastname1 = base64.b64decode(lastname_base64_decoded)
        lastname = lastname1.decode('UTF-8')
        print("Last name after decodingbytes: ", lastname)

        actual_email = content['email'].encode('UTF-8')
        print(actual_email)
        email_base64 = base64.b64encode(actual_email)
        print(email_base64)
        print(email_base64.decode('UTF-8'))
        email_base64_decoded = base64.b64decode(email_base64)
        email1 = base64.b64decode(email_base64_decoded)
        email = email1.decode('UTF-8')
        print("Email name after decodingbytes: ", email)
        password = '123455678'

    # FIRST WE WILL VERFY IF THE USER ALREADY THERE, IF YES, THEN JUST LOG HIM IN.
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""SELECT * FROM auth_user WHERE email = '%s' """ % email)
    result = cur.fetchall()
    if len(result) > 0:
        print('user already there, so login him in.')
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        #json_data = json.loads(json_data)
        # print(type(json_data))
        return logintoken.login_with_token_fb(json_data, app)
    else:
        active = 'TRUE'
        ts = time.time()
        doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
        #cipher_password = encrypt(encdcr_key, password)
        #plain_password = decrypt(encdcr_key, cipher_password)
        cipher_password = password
        cur.execute("INSERT INTO auth_user(first_name, last_name, email, password, active, created_at, login_from) VALUES(%s, %s, %s, %s, %s, %s, %s) RETURNING id;",
                    (firstname, lastname, email, cipher_password, active, doc, 'google'))
        conn.commit()
        user_id = cur.fetchone()[0]
        cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                    (user_id, 'preferred_language', 'us'))
        conn.commit()
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        conn.close()
        return logintoken.login_with_token_fb(json_data, app)
#########################################################
# Added by Haseeb
#########################################################


@blocked_bp.route('/api/signupinFacebook_BLOCKED', methods=['GET', 'POST'])
@blocked
def signupinFacebook_BLOCKED():
    """View function for signup view."""

    if 'firstname' in request.form and 'lastname' in request.form and 'email' in request.form:
        #print ("This is if part")
        actual_firstname = request.form['firstname'].encode('UTF-8')
        print(actual_firstname)
        print(type(actual_firstname))
        firstname_base64 = base64.b64encode(actual_firstname)
        print(firstname_base64)
        str_decoded = firstname_base64.decode()
        print(type(str_decoded))
        print(str_decoded)
        firstname_base64_decoded = base64.b64decode(firstname_base64)
        firstname1 = base64.b64decode(firstname_base64_decoded)
        firstname = firstname1.decode('UTF-8')
        print("First name after decodingbytes: ", firstname)

        actual_lastname = request.form['lastname'].encode('UTF-8')
        print(actual_lastname)
        lastname_base64 = base64.b64encode(actual_lastname)
        print(lastname_base64)
        print(lastname_base64.decode('UTF-8'))
        lastname_base64_decoded = base64.b64decode(lastname_base64)
        lastname1 = base64.b64decode(lastname_base64_decoded)
        lastname = lastname1.decode('UTF-8')
        print("Last name after decodingbytes: ", lastname)

        actual_email = request.form['email'].encode('UTF-8')
        print(actual_email)
        email_base64 = base64.b64encode(actual_email)
        print(email_base64)
        print(email_base64.decode('UTF-8'))
        email_base64_decoded = base64.b64decode(email_base64)
        email1 = base64.b64decode(email_base64_decoded)
        email = email1.decode('UTF-8')
        print("Email name after decodingbytes: ", email)
        password = '123455678'
    else:
        #print ("This is else part")
        content = request.get_json(silent=True)
        actual_firstname = content['firstname'].encode('UTF-8')
        print(actual_firstname)
        print(type(actual_firstname))
        firstname_base64 = base64.b64encode(actual_firstname)
        print(firstname_base64)
        str_decoded = firstname_base64.decode()
        print(type(str_decoded))
        print(str_decoded)
        firstname_base64_decoded = base64.b64decode(firstname_base64)
        firstname1 = base64.b64decode(firstname_base64_decoded)
        firstname = firstname1.decode('UTF-8')
        print("First name after decodingbytes: ", firstname)

        actual_lastname = content['lastname'].encode('UTF-8')
        print(actual_lastname)
        lastname_base64 = base64.b64encode(actual_lastname)
        print(lastname_base64)
        print(lastname_base64.decode('UTF-8'))
        lastname_base64_decoded = base64.b64decode(lastname_base64)
        lastname1 = base64.b64decode(lastname_base64_decoded)
        lastname = lastname1.decode('UTF-8')
        print("Last name after decodingbytes: ", lastname)

        actual_email = content['email'].encode('UTF-8')
        print(actual_email)
        email_base64 = base64.b64encode(actual_email)
        print(email_base64)
        print(email_base64.decode('UTF-8'))
        email_base64_decoded = base64.b64decode(email_base64)
        email1 = base64.b64decode(email_base64_decoded)
        email = email1.decode('UTF-8')
        print("Email name after decodingbytes: ", email)
        password = '123455678'
    # FIRST WE WILL VERFY IF THE USER ALREADY THERE, IF YES, THEN JUST LOG HIM IN.
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""SELECT * FROM auth_user WHERE email = '%s' """ % email)
    result = cur.fetchall()
    if len(result) > 0:
        print('user already there, so login him in.')
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        #json_data = json.loads(json_data)
        # print(type(json_data))
        return logintoken.login_with_token_fb(json_data, app)
    else:
        active = 'TRUE'
        ts = time.time()
        doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
        #cipher_password = encrypt(encdcr_key, password)
        #plain_password = decrypt(encdcr_key, cipher_password)
        cipher_password = password
        cur.execute("INSERT INTO auth_user(first_name, last_name, email, password, active, created_at, login_from) VALUES(%s, %s, %s, %s, %s, %s, %s) RETURNING id;",
                    (firstname, lastname, email, cipher_password, active, doc, 'facebook'))
        conn.commit()
        user_id = cur.fetchone()[0]
        cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                    (user_id, 'preferred_language', 'us'))
        conn.commit()
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        return logintoken.login_with_token_fb(json_data, app)


@blocked_bp.route('/api/signupFormProcess_BLOCKED', methods=['GET', 'POST'])
@blocked
def signupFormProcess_BLOCKED():
    """View function for signup view."""
    content = request.get_json(silent=True)
    print(content['name'])
    firstname = base64.b64decode(content['name']).decode('utf-8')
    lastname = ''
    email = base64.b64decode(content['email']).decode('utf-8')
    password = base64.b64decode(content['password']).decode('utf-8')
    # FIRST WE WILL VERFY IF THE USER ALREADY THERE, IF YES, THEN JUST LOG HIM IN.
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""SELECT * FROM auth_user WHERE email = '%s' """ % email)
    result = cur.fetchall()
    if len(result) > 0:
        print('user already there, so login him in.')
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        #json_data = json.loads(json_data)
        # print(type(json_data))
        """
        for key, value in json_data.items():
            print(key)
            print(value)
        print(json_data)
        """
        return logintoken.login_with_token_fb(json_data, app)
    else:
        active = 'TRUE'
        ts = time.time()
        doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
        #cipher_password = encrypt(encdcr_key, password)
        #plain_password = decrypt(encdcr_key, cipher_password)
        cipher_password = password
        cur.execute("INSERT INTO auth_user(first_name, last_name, email, password, active, created_at, login_from) VALUES(%s, %s, %s, %s, %s, %s, %s) RETURNING id;",
                    (firstname, lastname, email, cipher_password, active, doc, 'facebook'))
        conn.commit()
        user_id = cur.fetchone()[0]
        cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                    (user_id, 'preferred_language', 'us'))
        conn.commit()
        json_data = {}
        json_data['email'] = email
        json_data['password'] = password
        json_data['remember_me'] = ''
        json_data = {"remember_me": "", "email": email, "password": password}
        conn.close()
        return logintoken.login_with_token_fb(json_data, app)
###############################################################################################################################################


@blocked_bp.route('/api/signUpUserDotNet_BLOCKED', methods=['GET', 'POST'])
@blocked
def signUpUserDotNet_BLOCKED():
    """View function for signup view."""
    content = request.get_json(silent=True)
    if 1:
        firstname = content['firstname']
        lastname = content['lastname']
        email = content['email']
        password = content['password']
        active = 'TRUE'
        ts = time.time()
        doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
        #cipher_password = encrypt(encdcr_key, password)
        #plain_password = decrypt(encdcr_key, cipher_password)
        cipher_password = password
        # SELECT DATA FROM THE DATABASE
        conn = getConnection()
        cur = conn.cursor()
        cur.execute(
            """SELECT * FROM auth_user WHERE email = '%s' """ % email)
        result = cur.fetchall()
        if len(result) <= 0:
            cur.execute("INSERT INTO auth_user(first_name, last_name, email, password, active, created_at) VALUES(%s, %s, %s, %s, %s, %s) RETURNING id;",
                        (firstname, lastname, email, cipher_password, active, doc))
            conn.commit()
            user_id = cur.fetchone()[0]
            cur.execute("INSERT INTO user_meta(user_id, meta_name, meta_value) VALUES(%s, %s, %s) RETURNING id;",
                        (user_id, 'preferred_language', 'us'))
            conn.commit()
            code = 200
            response = dict()
            response['signup'] = 1
            response['signup_message'] = 'Account Created Successfully, Please Sign In to Continue'
            response['user_id'] = user_id
        else:
            code = 400
            response = dict()
            data = {
                'signup': '0', 'signup_message': 'Sorry! this account already exists', 'user_id': 0}
            response['errors'] = data
        conn.close()
        return jsonify(dict(meta=dict(code=code), response=response))
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- SIGN-IN SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#


@blocked_bp.route('/api/loginuser_BLOCKED', methods=['GET', 'POST'])
@blocked
def login_BLOCKED():
    # Currently using this code for login. This code is working fine.
    """View function for login view."""
    loglogger.info('Logged in user')
    content = request.get_json(silent=True)
    print(type(content))
    for key, value in content.items():
        print(key)
        print(value)
    print(request.json)
    return logintoken.login_with_token(request, app)

    # ==========================================================================

    # Implement few conditions in Login

    # """View function for login view."""
    # logger.info('Logged in user')

    # """Get data returned from the server."""

    # content = request.get_json(silent=True)
    # print(content)
    # print(request.json)

    # conn = psycopg2.connect(database="gcx_live", user="postgres", password="123456", host="localhost", port="5432")
    # cur = conn.cursor()

    # response = logintoken.login_with_token(request, app)
    # id = response['user']['id']
    # print("response is: ", response)

    # # The actual position of this code is here I am just testing the code that's why I have dragged above.
    # cur. execute("""SELECT login_count FROM auth_user WHERE id = '%s'""" %id)
    # result = cur.fetchall()
    # print(result)

    # login_count = result[0][0]

    # response['user']['login_count'] = login_count

    # return response
######################################################################
# PROPER NOUN = Capitalization Added by Haseeb
######################################################################


@blocked_bp.route('/api/capitalization_BLOCKED', methods=['GET', 'POST'])
@blocked
def capitalization_BLOCKED():
    with open('/var/www/properNounList/Proper_noun.txt') as f:
        for line in f:
            proper_noun = line.rstrip()
            print("Proper nouns are in smaller case:", proper_noun)
            # Here Limitized only the first letter and it will be capitalized, and then add back the rest of the string unchanged:
            capitalize_proper_noun = [word[0].capitalize(
            ) + word[1:] for word in proper_noun.split()]
            print("Proper nouns are in upper case:", capitalize_proper_noun)
    data = {'capitalization': '1'}
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
#   -Vocabulary :


@blocked_bp.route('/api/updateVocabulary_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def updateVocabulary_BLOCKED():
    """ Vocabulary """
    if 'user_id' in request.form and 'value' in request.form:
        user_id = request.form['user_id']
        new_value = request.form['value']
    else:
        content = request.get_json(silent=True)
        user_id = content['user_id']
        new_value = content['value']
    # NOW UPDATE IN THE USER-META TABLE AGAINST USER-ID,
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""UPDATE user_meta SET meta_value=%s WHERE meta_name = 'enable_vocabulary_enhancement' AND user_id=%s """, (new_value, user_id))
    conn.commit()
    # AND SELECT AGAIN DATA FROM THE DATABASE AND RETURN TO THE USER
    cur. execute(
        """SELECT meta_value FROM user_meta WHERE meta_name = 'enable_vocabulary_enhancement' AND user_id=%s """, user_id)
    rows = cur.fetchall()
    data = {'updateVocabulary': '1', 'contents': rows[0][0]}
    conn.close()
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
######################################################################
#   -Document Type :


@blocked_bp.route('/api/documentType_BLOCKED', methods=['GET', 'POST'])
@auth_token_required
@blocked
def documentType_BLOCKED():
    """ Document Type """
    data = {'documentType': '1'}
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- PROFILE UPDATION SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#
#################################################
#   Account :    Form Save & Resend Frsh Data (Change Password)
#   Added by HASEEB
#################################################
#   Profile :   Form Save & Resend Frsh Data   (Block Keywords)


@blocked_bp.route('/api/updateProfile_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def updateProfile_BLOCKED():
    """ Profile """
    if 'profile_id' in request.form and 'firstname' in request.form and 'lastname' in request.form and 'email' in request.form and 'password' in request.form:
        profile_id = request.form['profile_id']
        firstname = request.form['firstname']
        lastname = request.form['lastname']
        email = request.form['email']
        password = request.form['password']
    else:
        content = request.get_json(silent=True)
        profile_id = content['profile_id']
        firstname = content['firstname']
        lastname = content['lastname']
        email = content['email']
        password = content['password']
    #enc_password = utils.encrypt_password(password)
    enc_password = password
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""SELECT * FROM auth_user WHERE password = %s AND id = %s """,
                (enc_password, profile_id))
    print("""SELECT * FROM auth_user WHERE password = '%s' AND id = %s """,
          (enc_password, profile_id))
    result = cur.fetchall()
    if len(result) > 0:
        # NOW UPDATE IN THE USER-META TABLE AGAINST USER-ID,
        cur.execute("""UPDATE auth_user SET first_name = %s, last_name = %s, email = %s WHERE id =%s """,
                    (firstname, lastname, email, profile_id))
        conn.commit()
        # AND SELECT AGAIN DATA FROM THE DATABASE AND RETURN TO THE USER
        cur.execute(
            """SELECT first_name, last_name, email FROM auth_user WHERE id = %s """, profile_id)
        rows = cur.fetchall()
        updated_first_name = rows[0][0]
        updated_last_name = rows[0][1]
        updated_email = rows[0][2]
        if firstname == updated_first_name and lastname == updated_last_name and email == updated_email:
            data = {'updateProfile': '1', 'contents': '1'}
        else:
            data = {'updateProfile': '1', 'contents': '0'}
    else:
        data = {'updateProfile': enc_password, 'contents': '-1'}
    conn.close()
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
################################################################################
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- PLAGIARISM SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#
#   -Plagiarism :


@blocked_bp.route('/api/updatePlagiarism_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def updatePlagiarism_BLOCKED():
    """ Plagiarism """
    if 'user_id' in request.form and 'value' in request.form:
        user_id = request.form['user_id']
        new_value = request.form['value']
    else:
        content = request.get_json(silent=True)
        user_id = content['user_id']
        new_value = content['value']
    # NOW UPDATE IN THE USER-META TABLE AGAINST USER-ID,
    conn = getConnection()
    cur = conn.cursor()
    cur.execute("""UPDATE user_meta SET meta_value=%s WHERE meta_name = 'enable_plagiarism_checking' AND user_id=%s """, (new_value, user_id))
    conn.commit()
    # AND SELECT AGAIN DATA FROM THE DATABASE AND RETURN TO THE USER
    cur. execute("""SELECT meta_value FROM user_meta WHERE meta_name = 'enable_plagiarism_checking' AND user_id=%s """,
                 user_id)
    rows = cur.fetchall()
    data = {'updatePlagiarism': '1', 'contents': rows[0][0]}
    conn.close()
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- CHECK PLAGIARISM #
#----------------------------------------------------------------------------------------------------------------------------------------#
# This API will take data from user and send it to plag_url to check plagiarism


@blocked_bp.route('/api/checkPlagiarism_BLOCKED', methods=['GET', 'POST'])
@blocked
def checkPlagiarism_BLOCKED():
    if 'text' in request.form:
        text = request.form['text']
    else:
        content = request.get_json(silent=True)
        text = content['text']
        text = urllib.parse.unquote(text)
    values = {'text': text}
    print(values)
    text = json.dumps(text)
    print(text)
    plag_url = "https://api.plagiarismdetector.us/api/checkGrammarAPI"
    json_response = requests.post(plag_url, data=json.dumps(values))
    return Response(json_response, status=200, mimetype='application/json')
    #data_response = requests.post("https://api.plagiarismdetector.us/api/checkGrammarAPI", data=values)
    #print("response: ", data_response)

    #json_response = json.dumps(data, indent=4, sort_keys=True, default=str)
    # return Response(json_response, status=200, mimetype='application/json')
################################################################################
# This dummy API is just designed to get response
# Here I am taking str data and returning into Json response


@blocked_bp.route('/api/checkPlagiarismResponse_BLOCKED', methods=['GET', 'POST'])
@blocked
def checkPlagiarismResponse_BLOCKED():
    if 'text' in request.form:
        text = request.form['text']
    else:
        content = request.get_json(silent=True)
        text = content['text']
        text = urllib.parse.unquote(text)
    text = json.dumps(text)
    print(text)
    str = """[{webs: [{title: "U.S. Navy Diver Careers | Navy.com", headline: "",…}], unique: "false",…},…]
    0: {webs: [{title: "U.S. Navy Diver Careers | Navy.com", headline: "",…}], unique: "false",…}
    1: {webs: [{title: "U.S. Navy Diver Careers | Navy.com", headline: "",…}], unique: "false",…}
    2: {webs: [{title: "U.S. Navy Diver Careers | Navy.com", headline: "",…}], unique: "false",…}
    3: {webs: [{title: "U.S. Navy Diver Careers | Navy.com",…}], unique: "false",…}
    4: {webs: [{title: "Navy Diver in Belgrade, MT - Jobing.com",…},…], unique: "false",…}
    5: {webs: [{title: "U.S. Navy Diver Careers | Navy.com", headline: "",…}], unique: "false",…}"""
    values = {'text': text}
    print(values)
    json_response = json.dumps(str, values)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- SUBSCRIPTIONS SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#
###############################################################################################################################################
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- APPS. SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#
#   Apps :  Form Save & Resend Frsh Data   (Simly showing Our Apps)


@blocked_bp.route('/api/apps_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def apps_BLOCKED():
    """ Apps """
    data = {'apps': '1'}
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- UPGRADE TO PRO. SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#
#   -Premium :


@blocked_bp.route('/api/premium_BLOCKED', methods=['GET', 'POST'])
@blocked
# @auth_token_required
def premium_BLOCKED():
    """ Premium """
    data = {'premium': '1'}
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- GRAMMAR CHECKING SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#


@blocked_bp.route('/api/checkGrammar_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def checkGrammar_BLOCKED():
    """Get dummy data returned from the server."""
    tool = grammar_check.LanguageTool('en-US')
    text1 = "A sentence with a error in the Hitchhikers Guide tot he Galaxy"
    content = request.get_json(silent=True)
    text1 = content['text']
    matches = tool.check(text1)
    data = {'Heroes': [language_check.correct(text1, matches)],
            'string': 'gc',
            'description': '2',
            'precontext': '1',
            'type': '1',
            'url': '1',
            'suggestions': '4'
            }
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
######################################################################


@blocked_bp.route('/api/checkGrammarAPI_BLOCKED', methods=['GET', 'POST'])
@blocked
# @auth_token_required
def checkGrammarAPI_BLOCKED():
    language = request.form['language']
    text = request.form['text']
    url = 'http://3.219.35.115:8083/v2/check'
    values = {'language': language, 'text': text}
    user_agent = 'Mozilla/5.0 (Windows NT 6.1; Win64; x64)'
    headers = {'User-Agent': user_agent}
    data = urllib.parse.urlencode(values)
    data = data.encode('UTF-8')
    req = urllib.request.Request(url, data, headers)
    try:
        with urllib.request.urlopen(req) as response:
            the_page = response.read()
    except urllib.error.HTTPError as e:
        # Return code error (e.g. 404, 501, ...)
        # ...
        print('HTTPError: {}'.format(e.code))
    except urllib.error.URLError:
        # Not an HTTP-specific error (e.g. connection refused)
        # ... print(e.reason)
        # START THE JAVA SERVER AGAIN AND SEND REQUEST.
        command_run = "java -cp /var/www/gcxToolx/LanguageTool-4.0/languagetool-server.jar org.languagetool.server.HTTPServer --port 8083 --public --languageModel /var/languagetool_ngrams/"
        p = subprocess.Popen(command_run, shell=True)
        time.sleep(3)
        data = urllib.parse.urlencode(values)
        data = data.encode('UTF-8')
        req = urllib.request.Request(url, data, headers)
        with urllib.request.urlopen(req) as response:
            the_page = response.read()
        result = the_page
        result = result.decode('UTF-8')
        output = json.loads(result)
    else:
        # 200
        # ... print('good')
        result = the_page
        result = result.decode('UTF-8')
        output = json.loads(result)
    Software_Details = {'name': 'GrammarIn', 'version': '1.02', 'buildDate': '2018-03-01 00:00:00',
                        'apiVersion': '1.02', 'status': 'stable'}
    output['software'] = Software_Details
    json_response = json.dumps(output, indent=4, sort_keys=True, default=str)
    return Response(json_response, status=200, mimetype='application/json')
###############################################################################################################################################
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- SOME MISCELLANEOUS SECTION #
#----------------------------------------------------------------------------------------------------------------------------------------#


@blocked_bp.route('/api/testAPI_BLOCKED', methods=['GET', 'POST'])
@blocked
def testAPI_BLOCKED():
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT title, contents, doc_url  FROM user_documents WHERE doc_url='153926368920563' """)
    rows = cur.fetchall()
    title = rows[0][0]
    contents = rows[0][1]
    doc_id = rows[0][2]
    ##################################################
    # create an instance of a word document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    #doc.add_heading(title, 0)
    # add a paragraph and store
    # the object in a variable
    doc_para = doc.add_paragraph(contents)
    # now save the document to a location
    doc.save('/var/www/GCX/docx/GCX-'+doc_id+'.docx')
    output = ''
    conn.close()
    json_response = json.dumps(output)
    return Response(json_response, status=200, mimetype='application/json')
###############################################################################################################################################


@blocked_bp.route('/api/checkDocument_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def checkDocument_BLOCKED():
    """Get dummy data returned from the server."""
    tool = grammar_check.LanguageTool('en-US')
    text1 = "A sentence with a error in the Hitchhikers Guide tot he Galaxy"
    content = request.get_json(silent=True)
    text1 = content['text']
    matches = tool.check(text1)
    data = {'Heroes': [language_check.correct(text1, matches)],
            'string': 'gc',
            'description': '2',
            'precontext': '1',
            'type': '1',
            'url': '1',
            'suggestions': '4'
            }
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
######################################################################


@blocked_bp.route('/api/checkStatus_BLOCKED', methods=['GET', 'POST'])
@blocked
@auth_token_required
def checkStatus_BLOCKED():
    """Get dummy data returned from the server."""
    tool = language_tool.LanguageTool('en-US')
    text1 = "A sentence with a error in the Hitchhikers Guide tot he Galaxy"
    content = request.get_json(silent=True)
    text1 = content['text']
    matches = tool.check(text1)
    data = {'Heroes': [language_tool.correct(text1, matches)],
            'type': '1',
            'key': '1',
            'value': '1'
            }
    json_response = json.dumps(data)
    return Response(json_response, status=200, mimetype='application/json')
######################################################################


@blocked_bp.route('/api/checkIt_BLOCKED', methods=['POST'])
@blocked
def checkIt_BLOCKED():
    """Get data returned from the server."""
    if 'temp_id' in request.form:
        temp_id = request.form['temp_id']
    else:
        content = request.get_json(silent=True)
        temp_id = content['temp_id']
    conn = getConnection()
    cur = conn.cursor()
    # NOW WEWILL ACTIVATET HE USER RECORDS AND SEND THE DETAILS BACK TO THE CLIENT.
    cur.execute("UPDATE auth_user SET first_name = 'Admin1' WHERE id = 1;")
    conn.commit()
    cur. execute("SELECT * FROM auth_user;")
    rows = cur.fetchall()
    data = {'contents': rows[0][5]}
    conn.close()
    json_response = json.dumps(data, indent=4, sort_keys=True, default=str)
    return Response(json_response, status=200, mimetype='application/json')
#---#-------------------------------------------------------------------------------------------------------------------------------------#
#--- IP PORT AND OTHER SETTINGS #
#----------------------------------------------------------------------------------------------------------------------------------------#
###############################################################################################################################################


@blocked_bp.route('/api/Algorithms_BLOCKED', methods=['GET', 'POST'])
@blocked
def Algorithms_BLOCKED():
    start_time = time.time()
    language = request.form['language']
    text = request.form['text']
    if text.find('">') == -1 or text.find('@font-face') == -1:
        ts = time.time()
        rand = randint(1, 20)
        if language == 'en-US' or language == 'en-GB' or language == 'en-ZA' or language == 'en-CA' or language == 'en-AU' or language == 'en-NZ':
            url = 'http://35.168.113.150:8081/v2/check'
        else:
            url = 'http://3.219.35.115:8083/v2/check'
        # print(url)
        values = {'language': language, 'text': text}
        user_agent = 'Mozilla/5.0 (Windows NT 6.1; Win64; x64)'
        headers = {'User-Agent': user_agent}
        data = urllib.parse.urlencode(values)
        data = data.encode('UTF-8')
        req = urllib.request.Request(url, data, headers)
        try:
            with urllib.request.urlopen(req) as response:
                the_page = response.read()
        except urllib.error.HTTPError as e:
            # Return code error (e.g. 404, 501, ...)
            # ...
            print('HTTPError: {}'.format(e.code))
        except urllib.error.URLError:
            url = 'http://3.219.35.115:8083/v2/check'
            data = urllib.parse.urlencode(values)
            data = data.encode('UTF-8')
            req = urllib.request.Request(url, data, headers)
            with urllib.request.urlopen(req) as response:
                the_page = response.read()
            result = the_page
            result = result.decode('UTF-8')
            output = json.loads(result)
        else:
            # 200
            # ... print('good')
            result = the_page
            result = result.decode('UTF-8')
            output = json.loads(result)
        # print(output)
        #////////////////////////////////////////////////////////////////////////////#
        if len(output) > 0:
            output['matches'].sort(key=operator.itemgetter('offset'))
            for index in range(len(output['matches'])):
                offset = (output['matches'][index]['offset'])
                length = (output['matches'][index]['length'])
                suggestion_list_length = len(output['matches'][index]['replacements'])
                text_wrong = text[offset: offset + length]
                output['matches'][index]['wrong'] = text_wrong
                if suggestion_list_length == 1:
                    suggestion_length = len(output['matches'][index]['replacements'][0]['value'])
                    if suggestion_length == 0:
                        output['matches'][index] = {}
                if suggestion_list_length > 7:
                    output['matches'][index]['replacements'] = output['matches'][index]['replacements'][:7]
            output['matches'] = list(filter(None, output['matches']))
        else:
            output['matches'] = ''
        #////////////////////////////////////////////////////////////////////////////#
        Text_Details = {}
        Text_Details['characters'] = len(text)
        Text_Details['words'] = len(text.replace('\\', '').split())
        #words_array = text.split()
        #words_array_len = len(words_array)
        sentences = text.count('.') + text.count('!') + text.count('?')
        if sentences > 0:
            Text_Details['sentences'] = sentences
        elif len(text) > 10:
            Text_Details['sentences'] = 1
        else:
            Text_Details['sentences'] = 0
        Text_Details['reading_time'] = 0
        Text_Details['speaking_time'] = (Text_Details['words']*0.46)
        Text_Details['readability'] = '0%'
        output['text_details'] = Text_Details
        hiddenSpellError = 0
        hiddenSuggestion = 0
        hiddenGrammarError = 0
        for index in range(len(output['matches'])):
            rule_id = (output['matches'][index]['rule']['id'])
            issueType = (output['matches'][index]['rule']['issueType'])
            if rule_id.find("SPELLER_RULE") >= 0 or rule_id.find("MORFOLOGIK_RULE") >= 0 or rule_id == "HUNSPELL_NO_SUGGEST_RULE" or rule_id == "HUNSPELL_RULE" or rule_id == "FR_SPELLING_RULE":
                hiddenSpellError = hiddenSpellError + 1
            elif issueType == 'style' or issueType == 'locale-violation' or issueType == 'register':
                hiddenSuggestion = hiddenSuggestion + 1
            else:
                hiddenGrammarError = hiddenGrammarError + 1
        Errors_Details = {}
        Errors_Details['SpellError'] = hiddenSpellError
        Errors_Details['Suggestion'] = hiddenSuggestion
        Errors_Details['GrammarError'] = hiddenGrammarError
        output['errors'] = Errors_Details
    else:
        output = {}
        output['matches'] = ''
    #////////////////////////////////////////////////////////////////////////////#
    end_time = time.time()
    total_time = end_time - start_time
    #////////////////////////////////////////////////////////////////////////////#
    Software_Details = {}
    Software_Details['name'] = 'GrammarCheckerX'
    Software_Details['version'] = '4.6'
    Software_Details['buildDate'] = '2019-05-06 13:09:00'
    Software_Details['apiVersion'] = '6.0'
    Software_Details['status'] = 'stable'
    Software_Details['total_time'] = total_time

    if language == 'en-US' or language == 'en-GB' or language == 'en-ZA' or language == 'en-CA' or language == 'en-AU' or language == 'en-NZ':
        Software_Details['server'] = '34H-81'
    else:
        Software_Details['server'] = '34H-82'
    output['software'] = Software_Details
    json_response = json.dumps(output, indent=4, sort_keys=True, default=str)
    return Response(json_response, status=200, mimetype='application/json')
######################################################################


@blocked_bp.route('/api/myFunction_BLOCKED', methods=['GET', 'POST'])
@blocked
def myFunction_BLOCKED():
    print("My First Function In Python")
    return ''
