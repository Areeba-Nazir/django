import os
import re
import docx
import html
import time
import datetime
import requests
import urllib.parse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from random import randint
from urllib.request import urlretrieve
from rest_framework import permissions, status
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from .auth import bearer_auth
from .factory import getConnection
from .functions import prepareOutputWeb


################################################################################################
#   DOC :   Save, Resend, Grammar Checker, etc. ()
################################################################################################


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getDocumentsListing(request, *args, **kwargs):
    """ Doc """
    if 'profile_id' in request.data:
        user_id = request.data['profile_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT doc_url, doc_name, spell_errors, grammar_errors, suggestion_errors, doc_id, title, contents, dou FROM user_documents WHERE status = 'active' AND contents != '' AND doc_name != '' AND title != 'None' AND user_id='%s' ORDER BY dou DESC LIMIT 100 """ % user_id)
    rows = cur.fetchall()
    data = {'getDocumentsListing': '1', 'contents': rows}
    conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################

# ---#------------------------------------------------------------------------------------------#
# --- DOCUMENT SECTION #
# ----------------------------------------------------------------------------------------------#

# Un-Used
@api_view(['GET', 'POST'])
@bearer_auth
def getDocContentsPublic(request, *args, **kwargs):
    """ Doc """
    if 'docID' in request.data:
        doc_id = request.data['docID']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT * FROM user_documents WHERE doc_url= '%s' """ % doc_id)
    rows = cur.fetchall()
    data = {'updatePlagiarism': '1', 'contents': rows}
    conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################


@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getDocDataByProfileDocID(request, *args, **kwargs):
    """ Doc """
    if 'profile_id' in request.data and 'doc_id' in request.data:
        user_id = request.data['profile_id']
        doc_id = request.data['doc_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    cur.execute(
        """SELECT * FROM user_documents WHERE user_id=%s AND doc_url=%s """, (user_id, doc_id))
    rows = cur.fetchall()
    data = {'updatePlagiarism': '1', 'contents': rows}
    conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################

# Un-Used
@api_view(['POST'])
@bearer_auth
def getNewDocIDbyProfileID(request, *args, **kwargs):
    if 'profile_id' in request.data:
        user_id = request.data['profile_id']
        contents = request.data['homeText'] if 'homeText' in request.data else None
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    print("user_id ==>>>>", user_id)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    title = ''
    doc_name = 'Untitled document'
    if contents:
        contents = contents.replace("'", "\\'")
    else:
        contents = ""
    try:
        cur.execute(
            "INSERT INTO user_documents(user_id, doc, dou, doc_url, title, spell_errors, total_words, grammar_errors, suggestion_errors, score, doc_name, contents) VALUES(%s, '%s', '%s', '%s', '%s', %s, %s, %s, %s, %s, '%s', '%s') RETURNING doc_id " %
            (user_id, doc, doc, doc_url, title, spell_errors, total_words, grammar_errors, suggestion_errors, score,
             doc_name, contents))
        new_id = cur.fetchone()[0]
        data = {'new_id': new_id, 'new_doc_id': new_doc_id}
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################


@api_view(['POST'])
@bearer_auth
def getNewDocIDForGuestByContents(request, *args, **kwargs):
    if 'homeText' in request.data and 'profile_id' in request.data:
        contents = request.data['homeText']
        user_id = request.data['profile_id']
        # version = request.data['version']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    if user_id == '1':
        new_doc_id = 'edit' + str(timestamp) + str(random)
    else:
        new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    title = ''
    doc_name = 'Untitled document'
    if contents:
        contents = contents.replace("'", "\\'")
    else:
        contents = ""
    # print(contents)
    sql = "INSERT INTO user_documents(user_id, doc, dou, doc_url, title, contents, spell_errors, total_words, grammar_errors, suggestion_errors, score, doc_name) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id"
    try:
        cur.execute(sql, (user_id, doc, doc, doc_url, title, contents, spell_errors,
                          total_words, grammar_errors, suggestion_errors, score, doc_name))
        new_id = cur.fetchone()[0]
        data = {'new_id': new_id, 'new_doc_id': new_doc_id}
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################

# For upload of PDF Files
@api_view(['POST'])
def getNewDocIDbyProfileIDFileContents(request, *args, **kwargs):
    try:
        user_id = request.data['profile_id']
        file = request.data['file']
        create_chunks = request.data.get('create_chunks', 'false')
    except Exception as e:
        print(e)
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    timestamp = int(time.time())
    random = randint(100000, 999999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    # file.filename
    filename = new_doc_id + '.' + file.name.split('.')[-1]
    try:
        default_storage.save(os.path.join('/var/www/GCX/upload', filename), ContentFile(file.read()))
        #default_storage.save(os.path.join('/Work/GCX-Web/GCX-Django/gcx_django/uploads', filename), ContentFile(file.read()))
        #file.save(os.path.join('/var/www/GCX/upload', filename))
    except:
        for ch in ['#', '<', '$', '+', '%', '>', '!', '`', '&', '*', "'", '|', '{', '?', '\"', '=', '}', '/', ':',
                   '\\', '@']:
            if ch in filename:
                filename = filename.replace(ch, '')
        #file.save(os.path.join('/var/www/GCX/upload', filename))
        #default_storage.save(os.path.join('/Work/GCX-Web/GCX-Django/gcx_django/uploads', filename), ContentFile(file.read()))
        default_storage.save(os.path.join('/var/www/GCX/upload', filename), ContentFile(file.read()))
    try:
        if filename.split('.')[-1] == 'docx':
            doc = docx.Document(os.path.join('/var/www/GCX/upload', filename))
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            contents = '\n'.join(fullText)
            #text = contents.strip()
        else:
            url = "https://v2.pdfapis.com/api/v2/TextExtractor"
            files = {'file': open('/var/www/GCX/upload/' + filename, 'rb')}
            #files = {'file': open('/Work/GCX-Web/GCX-Django/gcx_django/uploads/' + filename, 'rb')}
            payload = {'authToken': '7FE9B6118D6484C8314F0378399FBD05', 'type': '3'}
            headers = {}
            r = requests.request("POST", url, headers=headers, data=payload, files=files)
            r = r.json()
            #contents = html.escape(r['text']).replace('\x00', '')
            contents = r['text']
            #text = re.sub("\r\n(\s)+", "\n", contents.strip())
            contents = re.sub("(\t)+", " ", contents)
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        json_response = {'message': 'Error reading file'}
        return Response(json_response, status=500)
    # print(contents)
    contents = re.sub("\r\n(\s)*", "\n", contents.strip())
    #contents = contents.replace("&quot;", '"').replace("&#x27;", "'").replace("&amp;", '&').replace("&gt;", '>').replace("&lt;", '<')
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()  # The point where time begins in the form of seconds.
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    sql = "INSERT INTO user_documents(user_id, doc, contents, doc_url, spell_errors, total_words, grammar_errors, suggestion_errors, score) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id"
    try:
        cur.execute(sql, (user_id, doc, contents, doc_url, spell_errors, total_words, grammar_errors, suggestion_errors, score))
        new_id = cur.fetchone()[0]
        # Chunk Slicing is working and is commented till further notice.
        #text = re.sub("\r\n(\s)*", "\n", r['text'].strip())
        chunklist = []
        if create_chunks.lower() == 'true':
            chnk_OBJ = chunkslicing()
            chnk_OBJ.checkConditions({'text': contents, 'length': len(contents), 'offset': 0, 'id': 0, })
            contents, chunklist = chnk_OBJ.makechunks()
        data = {'new_id': new_id, 'new_doc_id': new_doc_id, 'new_contents': contents, 'chunk_ids': chunklist}
        # Disabbled because creating double entries in database.
        #conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################

# Un-USed
@api_view(['POST'])
# @cross_origin(headers=['Content-Type']) # allow all origins all methods.
# @permission_classes([permissions.IsAuthenticated])
@bearer_auth
def getNewDocIDbyProfileIDContents(request, *args, **kwargs):
    if 'profile_id' in request.data and 'contents' in request.data:
        user_id = request.data['profile_id']
        contents = request.data['contents']
        contents = urllib.parse.unquote(contents)
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    # contents = contents)
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    contents = contents.replace("'", "\\'")
    contents = contents.replace("“", '"')
    contents = contents.replace("”", '"')
    contents = contents.replace("‘", "'")
    contents = contents.replace("’", "'")
    contents = contents.replace(",", ",")
    contents = contents.replace("„", ",")
    contents = contents.replace("…", ".")
    # print(contents)
    sql = "INSERT INTO user_documents(user_id, doc, contents, doc_url, spell_errors, total_words, grammar_errors, suggestion_errors, score) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id"
    try:
        cur.execute(sql, (user_id, doc, contents, doc_url, spell_errors,
                          total_words, grammar_errors, suggestion_errors, score))
        new_id = cur.fetchone()[0]
        data = {'new_id': new_id, 'new_doc_id': new_doc_id,
                'new_contents': contents, 'USER_ID': user_id}
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################


@api_view(['POST'])
# @permission_classes([permissions.IsAuthenticated])
# @bearer_auth
def getNewDocIDbyProfileIDContentsbyURL(request, *args, **kwargs):
    if 'profile_id' in request.data and 'file_path' in request.data and 'file_type' in request.data:
        user_id = request.data['profile_id']
        file_path = request.data['file_path']
        file_type = request.data['file_type']
        create_chunks = request.data.get('create_chunks', 'false')
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    if file_type == 'docx':
        randName = str(randint(100000, 999999))
        sampleFileName = "/var/www/GCX/docx/" + randName + ".docx"
        urlretrieve(file_path, sampleFileName)
        try:
            doc = docx.Document(sampleFileName)
        except Exception as e:
            print(e)
            # loglogger.error("Error Happend in execution.", exc_info=1)
            json_response = {"status": 501, "message": 'Error reading file'}
            code = 501
            return Response(json_response, code)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        contents = '\n'.join(fullText)
    elif file_type == 'txt':
        randName = str(randint(100000, 999999))
        sampleFileName = "/var/www/GCX/docx/" + randName + ".txt"
        urlretrieve(file_path, sampleFileName)
        fullText = []
        # The "r/rt" parameter in the open() function means "we're opening this file to read text data"
        with open(sampleFileName, "r") as f:
            lines = f.readlines()
            # Loop through all lines.
            for para in lines:
                fullText.append(para)
            contents = '\n'.join(fullText)
    else:
        randName = str(randint(100000, 999999))
        sampleFileName = "/var/www/GCX/docx/GCX-Report-" + randName + ".pdf"
        urlretrieve(file_path, sampleFileName)
        url = "http://v2.pdfapis.com/api/v2/TextExtractor"
        payload = {'authToken': '7FE9B6118D6484C8314F0378399FBD05', 'type': '3'}
        files = [('', ('.pdf', open(sampleFileName, 'rb'), 'application/pdf'))]
        headers = {}
        response = requests.request("POST", url, headers=headers, data=payload, files=files)
        #content = request.get_json(silent=True)
        r = response.json()
        try:
            contents = r['text'].replace('\x00', '')
            contents = re.sub("(\t)+", " ", contents)
        except Exception as e:
            print(e)
            # loglogger.error("Error Happend in execution.", exc_info=1)
            json_response = {"message": 'Error reading file'}
            code = 501
            return Response(json_response, code)
    contents = re.sub("\r\n(\s)*", "\n", contents.strip())
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    #contents = contents.replace("'", "\\'")
    contents = contents.replace("“", '"')
    contents = contents.replace("”", '"')
    contents = contents.replace("‘", "'")
    contents = contents.replace("’", "'")
    contents = contents.replace(",", ",")
    contents = contents.replace("„", ",")
    contents = contents.replace("…", ".")
    # print(contents)
    sql = "INSERT INTO user_documents(user_id, doc, contents, doc_url, spell_errors, total_words, grammar_errors, suggestion_errors, score) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id"
    try:
        cur.execute(sql, (user_id, doc, contents, doc_url, spell_errors,
                          total_words, grammar_errors, suggestion_errors, score))
        new_id = cur.fetchone()[0]
        # Chunk Slicing is working and is commented till further notice.
        chunklist = []
        if create_chunks.lower() == 'true':
            chnk_OBJ = chunkslicing()
            chnk_OBJ.checkConditions({'text': contents, 'length': len(contents), 'offset': 0, 'id': 0, })
            contents, chunklist = chnk_OBJ.makechunks()
        data = {'new_id': new_id, 'new_doc_id': new_doc_id, 'new_contents': contents, 'USER_ID': user_id, 'chunk_ids': chunklist}
    # Disabbled because creating double entries in database.
        #conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################


@api_view(['POST'])
@bearer_auth
def getNewDocIDbyProfileIDusingFileName(request, *args, **kwargs):
    if 'profile_id' in request.data and 'file_name' in request.data and 'docTitle' in request.data and 'contents' in request.data:
        user_id = request.data['profile_id']
        doc_name = request.data['file_name']
        title = request.data['docTitle']
        contents = request.data['contents']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    timestamp = int(time.time())
    random = randint(100000, 999999)
    new_doc_id = str(timestamp) + "" + str(user_id) + "" + str(random)
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    doc_url = new_doc_id
    spell_errors = 0
    total_words = 0
    grammar_errors = 0
    suggestion_errors = 0
    score = 0
    contents = contents.replace("'", "\\'")
    contents = contents.replace("“", '"')
    contents = contents.replace("”", '"')
    contents = contents.replace("‘", "'")
    contents = contents.replace("’", "'")
    contents = contents.replace(",", ",")
    contents = contents.replace("„", ",")
    contents = contents.replace("…", ".")
    # print(contents)
    sql = "INSERT INTO user_documents(user_id, doc, dou, doc_url, title, contents, spell_errors, total_words, grammar_errors, suggestion_errors, score, doc_name) VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING doc_id "
    try:
        cur.execute(sql, (user_id, doc, doc, doc_url, title, contents, spell_errors,
                          total_words, grammar_errors, suggestion_errors, score, doc_name))
        new_id = cur.fetchone()[0]
        data = {'new_id': new_id, 'new_doc_id': new_doc_id, 'USER_ID': user_id}
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################


@api_view(['GET', 'POST'])
@bearer_auth
def deleteDocumentByID(request, *args, **kwargs):
    if 'profile_id' in request.data and 'doc_id' in request.data:
        user_id = request.data['profile_id']
        doc_id = request.data['doc_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    try:
        cur.execute(
            """UPDATE user_documents SET status='deleted' WHERE doc_id = %s AND user_id = %s """, (doc_id, user_id))
        cur.execute(
            """SELECT status FROM user_documents WHERE doc_id = '%s' """ % doc_id)
        user_status = cur.fetchone()[0]
        data = {'deleteDocumentByID': '1', 'contents': user_status}
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Document Deletion failed'
        return Response(json_response, status=400)
    finally:
        conn.close()
    return Response(data, status.HTTP_200_OK)


################################################################################################
# Save File by Muneeb
################################################################################################


@api_view(['POST'])
@bearer_auth
def saveDocbyProfileIDusingDocID(request, *args, **kwargs):
    if 'profile_id' in request.data and 'file_name' in request.data and 'docTitle' in request.data and 'contents' in request.data and 'doc_id' in request.data:
        user_id = request.data['profile_id']
        doc_name = request.data['file_name']
        title = request.data['docTitle']
        contents = request.data['contents']
        doc_id = request.data['doc_id']
    else:
        error_response = {'response': 'Please Provide all parameters.', 'status': 'Failed'}
        return Response(error_response, status=status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    ts = time.time()
    doc = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    query = "UPDATE user_documents SET contents=%s, dou=%s, doc_name=%s, title=%s WHERE user_id=%s and doc_url=%s;"
    try:
        cur.execute(query, (contents, doc, doc_name, title, user_id, doc_id))
        conn.commit()
    except Exception as e:
        print(e)
        # loglogger.error("Error Happend in execution.", exc_info=1)
        conn.rollback()
        json_response = 'Saving data failed'
        conn.close()
        return Response(json_response, status=400)
    data = {"File saved successfully !"}
    conn.close()
    return Response(data, status.HTTP_200_OK)


@api_view(['POST'])
def getchunkslicing(request, *args, **kwargs):
    content = request.data
    try:
        text = content['text']
        language = 'en-US'
        # doc_id = content['doc_url']
    except Exception as e:
        print(e)
        error_response = {'response': 'Please Provide all parameters.', 'status': "Failed", 'code': "206"}
        return Response(error_response, status.HTTP_206_PARTIAL_CONTENT)
    threshold = 99
    chnk_OBJ = chunkslicing()
    chnk_OBJ.checkConditions({'text': text, 'length': len(text), 'offset': 0, 'id': 0,})
    #json_response = json.dumps(chnk_OBJ.getChunks(), indent=4, sort_keys=True, default=str)
    json_response = chnk_OBJ.getChunks()
    return Response(json_response, status.HTTP_200_OK)

@api_view(['POST'])
def createchunks(request, *args, **kwargs):
    content = request.data
    try:
        text = content['text']
        if 'doc_url' in content:
            doc_id = content['doc_url']
        else:
            doc_id = None
    except Exception as e:
        print(e)
        error_response = {'response': 'Please Provide all parameters.', 'status': "Failed", 'code': "206"}
        return Response(error_response, status.HTTP_206_PARTIAL_CONTENT)
    conn = getConnection()
    cur = conn.cursor()
    if doc_id:
        cur.execute("SELECT contents FROM user_documents WHERE doc_url = '%s' AND status = 'active';"% doc_id)
        row = cur.fetchone()
        if row:
            text = row[0]
    chnk_OBJ = chunkslicing()
    chnk_OBJ.checkConditions({'text': text, 'length': len(text), 'offset': 0, 'id': 0,})
    content, chunklist = chnk_OBJ.makechunks()
    conn.close()
    #json_response = json.dumps({'content': content, 'chunk_ids' : chunklist}, indent=4, sort_keys=True, default=str)
    return Response({'content': content, 'chunk_ids' : chunklist}, status.HTTP_200_OK)


class chunkslicing:
    def __init__(self):
        #self.chunk = {'text': content, 'length': len(content), 'offset': 0, 'id': 0,}
        self.characterLimit = 99
        self.abbrevations = []
        self.newchunks = []
        self.chunk_ids = []
        with open('/var/www/GCX/gcxAPIxDjango/gcx_django/gcxAPIx/abbrevations.txt', 'r') as fileReader:
            for item in fileReader.readlines():
                self.abbrevations.append(item.strip('\n'))

    def checkConditions(self, chunk):
        text = chunk['text']
        matches = re.finditer('\s[A-Za-z]+(\)*[.](\s)+[A-Za-z]+)\s', text)
        ci = 0
        id = chunk['id']
        for match in matches:
            if match.start() < self.characterLimit:
                continue
            elif match.end() != chunk['length']:
                dot_index = match.group().find('.') + 1
                ci = match.start() + dot_index
                if text[match.start():ci].strip().lower() in self.abbrevations:
                    continue
                if text[ci] == '\n':
                    ci += 1
                new_text = text[:ci]
                self.newchunks.append({'id': str(id) + '-A', 'text': new_text, 'offset': chunk['offset'], 'length': ci})
                self.chunk_ids.append(str(id) + '-A')
                id += 1
                self.checkConditions({'id': id, 'text': text[ci:], 'offset': chunk['offset']+ci, 'length': chunk['length']-ci})
                return
            else:
                break
        self.newchunks.append({'id': str(id) + '-A', 'text': text[ci:], 'offset': chunk['offset'], 'length': chunk['length']})
        self.chunk_ids.append(str(id) + '-A')
        return 0

    def getChunks(self):
        return self.newchunks

    def makechunks(self):
        content = ""
        for chnk in self.newchunks:
            content += "<chunk id={} style='/*white-space: pre-wrap*/'>{}</chunk>".format(chnk["id"], chnk["text"])
        return content, self.chunk_ids

@api_view(['GET'])
@bearer_auth
def getAbbreviations(request, *args, **kwargs):
    abbrevations = list()
    with open('/var/www/GCX/gcxAPIxDjango/gcx_django/gcxAPIx/abbrevations.txt', 'r') as fileReader:
        for item in fileReader.readlines():
            abbrevations.append(item.strip('\n'))
    json_response = {'list': abbrevations}
    return Response(json_response, status.HTTP_200_OK)